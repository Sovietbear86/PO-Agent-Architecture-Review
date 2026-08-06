from pathlib import Path
from typing import Any

from pypdf import PdfReader
from openpyxl import load_workbook
from docx import Document
import extract_msg


class AttachmentReader:
    def read(self, path: Path) -> dict[str, Any]:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._read_pdf(path)
        if suffix in {".xlsx", ".xlsm"}:
            return self._read_xlsx(path)
        if suffix == ".docx":
            return self._read_docx(path)
        if suffix == ".msg":
            return self._read_msg(path)
        if suffix in {".txt", ".md", ".csv", ".json", ".xml", ".yaml", ".yml"}:
            return {"type": suffix[1:], "text": path.read_text(encoding="utf-8", errors="replace")}
        raise ValueError(f"Неподдерживаемый формат: {suffix}")

    def _read_pdf(self, path: Path) -> dict[str, Any]:
        reader = PdfReader(str(path))
        pages = [
            {"page": idx + 1, "text": page.extract_text() or ""}
            for idx, page in enumerate(reader.pages)
        ]
        return {"type": "pdf", "pages": pages}

    def _read_xlsx(self, path: Path) -> dict[str, Any]:
        wb = load_workbook(path, read_only=True, data_only=False)
        sheets: list[dict[str, Any]] = []
        for ws in wb.worksheets:
            rows = [[cell.value for cell in row] for row in ws.iter_rows()]
            sheets.append({"name": ws.title, "rows": rows})
        return {"type": "xlsx", "sheets": sheets}

    def _read_docx(self, path: Path) -> dict[str, Any]:
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        tables = [
            [[cell.text for cell in row.cells] for row in table.rows]
            for table in doc.tables
        ]
        return {"type": "docx", "paragraphs": paragraphs, "tables": tables}

    def _read_msg(self, path: Path) -> dict[str, Any]:
        msg = extract_msg.Message(str(path))
        return {
            "type": "msg",
            "subject": msg.subject,
            "sender": msg.sender,
            "to": msg.to,
            "date": str(msg.date) if msg.date else None,
            "body": msg.body,
        }
